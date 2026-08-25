#!/usr/bin/env python3
"""Build a final SAD .docx by cloning the company template (for styles/cover/TOC) and filling in the
generated Markdown body. Diagrams must already be PNGs referenced via ![](path) — run
render_diagrams.py first.

Strategy: keep the template's front matter (cover page, revision tables, TOC) up to its first
Heading-1, drop the template's *example* body, then append the generated content using the template's
own styles (Heading 1..6, Normal, Table Grid). This preserves the official look without us re-creating
styles. Markdown supported: ATX headings (#..######), GFM pipe tables, '- '/'* ' bullets, blockquotes,
images ![alt](path), and plain paragraphs. Inline **bold**/`code` markers are stripped to plain text.

The AU-QR-R&D-032 architecture template ships as a legacy binary .doc; python-docx can't open that.
If --template points at a .doc, this script first converts it to .docx via LibreOffice (soffice).

Usage:
  python3 build_sad_docx.py <sad.rendered.md> --template <template.doc|.docx> --out <out.docx> \
      [--img-base <dir>] [--keep-front-until "引言"]
"""
import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile

import docx
from docx.shared import Inches


def ensure_docx_template(path):
    """python-docx only opens OOXML .docx. If given a legacy .doc, convert it via soffice and return
    the converted path; otherwise return the path unchanged."""
    if not path.lower().endswith(".doc"):
        return path
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        print("ERROR: 模板是 .doc 且系统无 soffice/libreoffice 可转换；请先手动转成 .docx 再传入",
              file=sys.stderr)
        sys.exit(1)
    outdir = tempfile.mkdtemp(prefix="sad_tmpl_")
    r = subprocess.run([soffice, "--headless", "--convert-to", "docx", "--outdir", outdir, path],
                       capture_output=True, text=True)
    converted = os.path.join(outdir, os.path.splitext(os.path.basename(path))[0] + ".docx")
    if r.returncode != 0 or not os.path.exists(converted):
        print(f"ERROR: .doc → .docx 转换失败: {r.stderr[:300]}", file=sys.stderr)
        sys.exit(1)
    print(f"模板 .doc 已转换为 .docx：{converted}")
    return converted


NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def suppress_heading_autonumber(paragraph):
    """Cancel the Heading style's built-in auto-numbering for this paragraph.

    The company templates define Heading 1/2/3 WITH a multilevel auto-number (numId). Our markdown
    already writes the section numbers into the heading text ("4.1 功能 制氧控制"), so without this Word
    paints its own number on top — and since the doc title is a Heading 1, the auto-numbers come out
    offset by one, giving doubled/mismatched numbers. We author the numbers, so override each heading
    to numId=0 (no list). numPr must follow pStyle in pPr."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:numPr")):
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "0")
    numPr.append(ilvl)
    numPr.append(numId)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        pStyle.addnext(numPr)
    else:
        pPr.insert(0, numPr)


def style_id_to_name(doc):
    """styleId -> display name. The template's Heading-1 styleId is numeric (e.g. '2'), so we must
    resolve through styles.xml rather than guessing from the raw w:val."""
    mapping = {}
    for s in doc.styles:
        try:
            mapping[s.style_id] = s.name
        except Exception:
            pass
    return mapping


def first_heading1_index(doc, keyword=None):
    """Return the body-child index of the first Heading-1 paragraph (optionally matching keyword)."""
    id2name = style_id_to_name(doc)
    body = list(doc.element.body)
    for i, child in enumerate(body):
        if not child.tag.endswith("}p"):
            continue
        pPr = child.find(NS + "pPr")
        if pPr is None:
            continue
        pStyle = pPr.find(NS + "pStyle")
        if pStyle is None:
            continue
        val = pStyle.get(NS + "val") or ""
        name = id2name.get(val, val)
        if name == "Heading 1":
            if keyword:
                txt = "".join(t.text or "" for t in child.iter(NS + "t"))
                if keyword not in txt:
                    continue
            return i
    return None


def strip_inline(s):
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"\*(.+?)\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    return s.strip()


def style_or_default(doc, name, fallback="Normal"):
    """Resolve a paragraph/table style by its display NAME and return the style object.

    python-docx indexes doc.styles by style_id (e.g. 'Heading1'), NOT by the UI name ('Heading 1'),
    so a direct doc.styles['Heading 1'] raises KeyError and the body silently falls back to Normal —
    losing all heading structure/TOC. We map name -> style object ourselves. Returns the fallback
    style object (or None) when the requested name isn't present."""
    by_name = {}
    for s in doc.styles:
        try:
            by_name[s.name] = s
        except Exception:
            pass
    if name in by_name:
        return by_name[name]
    if fallback and fallback in by_name:
        return by_name[fallback]
    return None


def add_table(doc, header, rows, img_base, out_dir):
    ncol = len(header)
    t = doc.add_table(rows=1, cols=ncol)
    t.style = style_or_default(doc, "Table Grid", None) or t.style
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = strip_inline(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j in range(ncol):
            cells[j].text = strip_inline(row[j]) if j < len(row) else ""
    doc.add_paragraph("")


def parse_table_block(lines, start):
    """Parse a GFM pipe table starting at lines[start]. Return (header, rows, next_index)."""
    def split_row(l):
        l = l.strip()
        if l.startswith("|"):
            l = l[1:]
        if l.endswith("|"):
            l = l[:-1]
        return [c.strip() for c in l.split("|")]

    header = split_row(lines[start])
    i = start + 1
    # optional separator row of ---
    if i < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i]) and "-" in lines[i]:
        i += 1
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(split_row(lines[i]))
        i += 1
    return header, rows, i


def render_markdown(doc, md_text, img_base, out_dir):
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = strip_inline(m.group(2))
            p = doc.add_paragraph(text, style=style_or_default(doc, f"Heading {min(level,6)}"))
            suppress_heading_autonumber(p)  # keep our authored numbers; cancel the style's auto-number
            i += 1
            continue

        # image
        mi = re.match(r"^!\[[^\]]*\]\(([^)]+)\)\s*$", stripped)
        if mi:
            path = mi.group(1)
            if not os.path.isabs(path):
                cand = os.path.join(img_base, path) if img_base else path
                path = cand if os.path.exists(cand) else os.path.join(out_dir, path)
            if os.path.exists(path):
                try:
                    doc.add_picture(path, width=Inches(6.0))
                except Exception as e:
                    doc.add_paragraph(f"[图片插入失败 {path}: {e}]")
            else:
                doc.add_paragraph(f"[缺图片 {path}]")
            i += 1
            continue

        # table
        if stripped.startswith("|"):
            header, rows, ni = parse_table_block(lines, i)
            add_table(doc, header, rows, img_base, out_dir)
            i = ni
            continue

        # bullet
        mb = re.match(r"^[-*]\s+(.*)$", stripped)
        if mb:
            doc.add_paragraph(strip_inline(mb.group(1)),
                              style=style_or_default(doc, "List Bullet"))
            i += 1
            continue

        # blockquote / note
        if stripped.startswith(">"):
            doc.add_paragraph(strip_inline(stripped.lstrip(">").strip()))
            i += 1
            continue

        # plain paragraph
        doc.add_paragraph(strip_inline(stripped))
        i += 1


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md")
    ap.add_argument("--template", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--img-base", default=None,
                    help="base dir for relative image paths (default: md's dir)")
    ap.add_argument("--keep-front-until", default=None,
                    help="keep template front matter up to the first Heading-1 containing this "
                         "keyword (default: first Heading-1 of any text)")
    args = ap.parse_args()

    template = ensure_docx_template(args.template)
    doc = docx.Document(template)
    body = doc.element.body

    # remove template's example body (from first Heading-1 onward)
    idx = first_heading1_index(doc, args.keep_front_until)
    if idx is not None:
        children = list(body)
        for child in children[idx:]:
            # keep the final sectPr (page setup) if present
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)
    else:
        print("WARN: 模板中未找到 Heading 1 锚点，将直接在文末追加内容", file=sys.stderr)

    with open(args.md, encoding="utf-8") as f:
        md_text = f.read()
    img_base = args.img_base or (os.path.dirname(os.path.abspath(args.md)) or ".")
    out_dir = os.path.dirname(os.path.abspath(args.out)) or "."
    render_markdown(doc, md_text, img_base, out_dir)

    doc.save(args.out)
    print(f"已生成：{args.out}")


if __name__ == "__main__":
    main()
