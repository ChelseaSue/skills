#!/usr/bin/env python3
"""Dump a .docx into structured, readable text (headings + paragraphs + tables) in body order.

Why: the SRS workflow needs to *read* the system-requirements / template docx reliably rather than
eyeball the original Word file. This walks the document body in order so headings, normal text and
tables stay in their real positions, and renders tables as markdown-ish grids (de-duping merged cells).

Usage:
  python3 extract_docx.py <file.docx>                 # full structured dump
  python3 extract_docx.py <file.docx> --headings      # only the heading outline
  python3 extract_docx.py <file.docx> --tables        # only tables
  python3 extract_docx.py <file.docx> --max-rows 30   # cap rows printed per table (default: all)
  python3 extract_docx.py <file.docx> --extract-images <outdir>
        # save every embedded image to <outdir>, each named "<NN>_<heading>.png", and print which
        # heading each image sits under — so you can pick the real 系统框图/框图/外形图 and reference
        # it in the SRS markdown as ![](path). Later build_srs_docx.py embeds those PNGs into the docx.
"""
import argparse
import os
import re
import sys

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_blocks(document):
    """Yield ('p', Paragraph) / ('tbl', Table) in document body order."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield "p", Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield "tbl", Table(child, document)


def is_heading(style_name):
    if not style_name:
        return False
    s = style_name
    return "Heading" in s or "heading" in s or s.startswith("标题") or s.startswith("Title")


def render_table(table, max_rows=None):
    lines = []
    rows = table.rows
    n = len(rows) if max_rows is None else min(max_rows, len(rows))
    for r in rows[:n]:
        cells = [c.text.strip().replace("\n", " / ") for c in r.cells]
        # collapse consecutive duplicates produced by horizontal cell merges
        uniq = []
        for c in cells:
            if not uniq or uniq[-1] != c:
                uniq.append(c)
        lines.append("| " + " | ".join(uniq) + " |")
    if max_rows is not None and len(rows) > max_rows:
        lines.append(f"| ...（共 {len(rows)} 行，省略 {len(rows) - max_rows} 行）|")
    return "\n".join(lines)


def safe_name(s, maxlen=24):
    s = re.sub(r"[\s/\\:*?\"<>|]+", "_", s.strip())
    return s[:maxlen] or "img"


def extract_images(d, outdir):
    """Save embedded images to outdir, named by the heading they sit under, and report mapping."""
    os.makedirs(outdir, exist_ok=True)
    img_rels = {rid: r for rid, r in d.part.rels.items() if "image" in r.reltype}
    cur_h = ""
    n = 0
    for p in d.paragraphs:
        style = p.style.name if p.style else ""
        if is_heading(style) and p.text.strip():
            cur_h = p.text.strip()
        for blip in p._p.findall(".//" + qn("a:blip")):
            rid = blip.get(qn("r:embed"))
            rel = img_rels.get(rid)
            if rel is None:
                continue
            n += 1
            ext = os.path.splitext(rel.target_ref)[1] or ".png"
            fname = f"{n:02d}_{safe_name(cur_h)}{ext}"
            out = os.path.join(outdir, fname)
            with open(out, "wb") as f:
                f.write(rel.target_part.blob)
            print(f"{n:2d} 章节[{cur_h[:30]}] -> {out}")
    if n == 0:
        print("（文档中未找到内嵌图片）")
    else:
        print(f"\n共导出 {n} 张图片到 {outdir}/。挑出需要的（如系统框图/外形图），在 SRS 的 md 里用 ![](相对路径) 引用。")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("file")
    ap.add_argument("--headings", action="store_true", help="only print heading outline")
    ap.add_argument("--tables", action="store_true", help="only print tables")
    ap.add_argument("--max-rows", type=int, default=None, help="cap rows per table")
    ap.add_argument("--extract-images", metavar="OUTDIR",
                    help="save embedded images to OUTDIR, named by their heading")
    args = ap.parse_args()

    try:
        d = docx.Document(args.file)
    except Exception as e:
        print(f"ERROR opening {args.file}: {e}", file=sys.stderr)
        sys.exit(1)

    if args.extract_images:
        extract_images(d, args.extract_images)
        return

    tbl_idx = 0
    for kind, blk in iter_blocks(d):
        if kind == "p":
            if args.tables:
                continue
            text = blk.text.strip()
            if not text:
                continue
            style = blk.style.name if blk.style else ""
            if is_heading(style):
                # crude level from "Heading N"
                lvl = "".join(ch for ch in style if ch.isdigit()) or "1"
                print(f"\n{'#' * min(int(lvl), 6)} {text}")
            elif not args.headings:
                print(text)
        else:  # table
            if args.headings:
                tbl_idx += 1
                continue
            print(f"\n[表 #{tbl_idx} {len(blk.rows)}x{len(blk.rows[0].cells) if blk.rows else 0}]")
            print(render_table(blk, args.max_rows))
            tbl_idx += 1


if __name__ == "__main__":
    main()
