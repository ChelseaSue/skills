#!/usr/bin/env python3
"""
Convert SWDD Markdown files to DOCX, replacing mermaid/plantuml code blocks with PNG images.

Usage:
  python md_to_docx.py --swdd-root ./swdd [module_name ...]
  python md_to_docx.py --swdd-root ./swdd              # convert all modules
  python md_to_docx.py --swdd-root ./swdd ABBSM ADESC  # convert specific modules
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


IMG_WIDTH_INCHES = 6.0  # Max image width


def setup_heading_numbering(doc):
    """Set up multi-level numbering for Heading 2/3/4 styles.

    Creates a Word numbering definition that produces:
      Heading 2: 1, 2, 3 ...
      Heading 3: 1.1, 1.2, 2.1 ...
      Heading 4: 1.1.1, 2.4.1, 2.7.1 ...
    Heading 1 is left unnumbered (used for document title).
    """
    # Access or create numbering part
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part.numbering_definitions._numbering

    # Find a safe abstractNumId (avoid collision with existing ones)
    existing_ids = [int(e.get(qn('w:abstractNumId')))
                    for e in numbering_elm.findall(qn('w:abstractNum'))]
    abstract_num_id = max(existing_ids, default=-1) + 1

    existing_num_ids = [int(e.get(qn('w:numId')))
                        for e in numbering_elm.findall(qn('w:num'))]
    num_id = max(existing_num_ids, default=0) + 1

    # Build abstractNum XML with 3 levels (ilvl 0=Heading2, 1=Heading3, 2=Heading4)
    abstract_num_xml = f'''
    <w:abstractNum w:abstractNumId="{abstract_num_id}" {nsdecls("w")}>
      <w:multiLevelType w:val="multilevel"/>
      <w:lvl w:ilvl="0">
        <w:start w:val="1"/>
        <w:numFmt w:val="decimal"/>
        <w:lvlText w:val="%1"/>
        <w:lvlJc w:val="left"/>
        <w:pPr>
          <w:ind w:left="0" w:firstLine="0"/>
        </w:pPr>
      </w:lvl>
      <w:lvl w:ilvl="1">
        <w:start w:val="1"/>
        <w:numFmt w:val="decimal"/>
        <w:lvlText w:val="%1.%2"/>
        <w:lvlJc w:val="left"/>
        <w:pPr>
          <w:ind w:left="0" w:firstLine="0"/>
        </w:pPr>
      </w:lvl>
      <w:lvl w:ilvl="2">
        <w:start w:val="1"/>
        <w:numFmt w:val="decimal"/>
        <w:lvlText w:val="%1.%2.%3"/>
        <w:lvlJc w:val="left"/>
        <w:pPr>
          <w:ind w:left="0" w:firstLine="0"/>
        </w:pPr>
      </w:lvl>
    </w:abstractNum>
    '''.strip()

    num_xml = f'''
    <w:num w:numId="{num_id}" {nsdecls("w")}>
      <w:abstractNumId w:val="{abstract_num_id}"/>
    </w:num>
    '''.strip()

    numbering_elm.append(parse_xml(abstract_num_xml))
    numbering_elm.append(parse_xml(num_xml))

    return num_id


def strip_heading_number(heading_text):
    """Strip manual numbering prefix from heading text.

    Examples:
      '1 Overview' -> 'Overview'
      '1.1 Purpose' -> 'Purpose'
      '2.7.1 ABBSM_vidMainFunction' -> 'ABBSM_vidMainFunction'
    Returns (stripped_text, had_number).
    """
    m = re.match(r'^(\d+(?:\.\d+)*)\s+(.*)', heading_text)
    if m:
        return m.group(2), True
    return heading_text, False


def apply_numbering_to_paragraph(para, num_id, ilvl):
    """Apply multi-level numbering to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="{ilvl}"/>'
        f'  <w:numId w:val="{num_id}"/>'
        f'</w:numPr>'
    )
    pPr.insert(0, numPr)


def find_matching_png(module_dir, block_type, block_content, block_context_before):
    """Find the matching PNG for a mermaid/plantuml code block."""
    img_dir = module_dir / "img"
    if not img_dir.exists():
        return None

    module_name = module_dir.name

    # Strategy 1: Check if it's a Static Diagram (flowchart LR or contains "subgraph")
    if block_type == "mermaid" and ("flowchart LR" in block_content or "flowchart RL" in block_content):
        png = img_dir / f"{module_name}_Static_Diagram.png"
        if png.exists():
            return png

    # Strategy 2: Check if it's a Dynamic Behavior (plantuml/puml)
    if block_type in ("plantuml", "puml"):
        png = img_dir / f"{module_name}_Dynamic_Behavior.png"
        if png.exists():
            return png

    # Strategy 3: Look at context before the block for function name
    # Use findall + take last match to get the NEAREST heading (not the first/farthest)
    func_matches = re.findall(r'####\s+[\d.]+\s+(\w+)\s*$', block_context_before, re.MULTILINE)
    if func_matches:
        func_name = func_matches[-1]
        png = img_dir / f"{module_name}_{func_name}_Flowchart.png"
        if png.exists():
            return png
        png = img_dir / f"{func_name}_Flowchart.png"
        if png.exists():
            return png

    # Strategy 4: For mermaid flowchart TD blocks, find function name from context
    if block_type == "mermaid" and "flowchart TD" in block_content:
        for line in block_context_before.split('\n')[::-1]:
            m = re.search(r'####\s+[\d.]+\s+(\w+)', line)
            if m:
                func_name = m.group(1)
                png = img_dir / f"{module_name}_{func_name}_Flowchart.png"
                if png.exists():
                    return png
                png = img_dir / f"{func_name}_Flowchart.png"
                if png.exists():
                    return png
                break

    # Strategy 5: Static diagram with flowchart TB
    if block_type == "mermaid" and "Static Diagram" in block_context_before[-500:]:
        png = img_dir / f"{module_name}_Static_Diagram.png"
        if png.exists():
            return png

    return None


def parse_markdown_to_blocks(md_text, module_dir):
    """Parse markdown into a list of blocks: text, image, table, etc."""
    blocks = []
    lines = md_text.split('\n')
    i = 0
    text_buffer = []

    while i < len(lines):
        line = lines[i]

        # Check for code block start
        code_match = re.match(r'^```(mermaid|plantuml|puml)\s*$', line.strip())
        if code_match:
            if text_buffer:
                blocks.append(('text', '\n'.join(text_buffer)))
                text_buffer = []

            block_type = code_match.group(1)
            code_lines = []
            context_before = '\n'.join(lines[max(0, i-60):i])
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1

            block_content = '\n'.join(code_lines)
            png = find_matching_png(module_dir, block_type, block_content, context_before)
            if png:
                blocks.append(('image', str(png)))
            else:
                blocks.append(('code', block_content))
            continue

        # Check for other code blocks (bash, etc)
        other_code_match = re.match(r'^```(\w*)\s*$', line.strip())
        if other_code_match:
            if text_buffer:
                blocks.append(('text', '\n'.join(text_buffer)))
                text_buffer = []
            code_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            blocks.append(('code', '\n'.join(code_lines)))
            continue

        text_buffer.append(line)
        i += 1

    if text_buffer:
        blocks.append(('text', '\n'.join(text_buffer)))

    return blocks


def add_table_to_doc(doc, table_lines):
    """Parse markdown table and add to document."""
    data_lines = []
    for line in table_lines:
        stripped = line.strip()
        if stripped and not re.match(r'^\|[\s\-:|]+$', stripped):
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            data_lines.append(cells)

    if len(data_lines) < 1:
        return

    num_cols = max(len(row) for row in data_lines)
    for row in data_lines:
        while len(row) < num_cols:
            row.append('')

    table = doc.add_table(rows=len(data_lines), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(data_lines):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            para = cell.paragraphs[0]
            para.style = doc.styles['Normal']
            for run in para.runs:
                run.font.size = Pt(9)

            if r_idx == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for run in para.runs:
                    run.bold = True


def process_text_block(doc, text, num_id=None):
    """Process a text block, handling headings, tables, bold, italic, etc."""
    lines = text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)

            if num_id and level >= 2:
                heading_text, had_number = strip_heading_number(heading_text)
                para = doc.add_heading(heading_text, level=min(level, 4))
                if had_number:
                    apply_numbering_to_paragraph(para, num_id, level - 2)
            else:
                doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # Table detection
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_to_doc(doc, table_lines)
            continue

        # Figure caption
        fig_match = re.match(r'^\*{1,2}(Figure.*?)\*{1,2}$', stripped)
        if fig_match:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(fig_match.group(1))
            run.bold = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        add_formatted_text(para, stripped)
        i += 1


def add_formatted_text(para, text):
    """Add text with basic markdown formatting (bold, italic, code)."""
    pattern = re.compile(r'(\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)')
    pos = 0

    for match in pattern.finditer(text):
        if match.start() > pos:
            run = para.add_run(text[pos:match.start()])
            run.font.size = Pt(10)

        matched = match.group(0)
        if matched.startswith('**') and matched.endswith('**'):
            run = para.add_run(matched[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        elif matched.startswith('`') and matched.endswith('`'):
            run = para.add_run(matched[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif matched.startswith('*') and matched.endswith('*'):
            run = para.add_run(matched[1:-1])
            run.italic = True
            run.font.size = Pt(10)

        pos = match.end()

    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.size = Pt(10)


def add_image_to_doc(doc, img_path):
    """Add an image centered in the document."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()

    from PIL import Image as PILImage
    try:
        with PILImage.open(img_path) as img:
            w, h = img.size
            max_width = IMG_WIDTH_INCHES
            max_height = 8.0
            dpi = 96
            img_width_in = w / dpi
            img_height_in = h / dpi

            if img_width_in > max_width:
                scale = max_width / img_width_in
                img_width_in = max_width
                img_height_in *= scale

            if img_height_in > max_height:
                scale = max_height / img_height_in
                img_height_in = max_height
                img_width_in *= scale

            run.add_picture(img_path, width=Inches(img_width_in))
    except ImportError:
        run.add_picture(img_path, width=Inches(IMG_WIDTH_INCHES))
    except Exception as e:
        print(f"  Warning: Could not add image {img_path}: {e}")
        run.add_text(f"[Image: {os.path.basename(img_path)}]")


def convert_module(module_dir):
    """Convert a single module's MD to DOCX."""
    md_files = list(module_dir.glob("*.md"))
    if not md_files:
        return None

    md_file = md_files[0]
    module_name = module_dir.name
    docx_path = module_dir / md_file.name.replace('.md', '.docx')

    print(f"Converting {module_name}: {md_file.name} -> {docx_path.name}")

    md_text = md_file.read_text(encoding='utf-8')
    blocks = parse_markdown_to_blocks(md_text, module_dir)

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    num_id = setup_heading_numbering(doc)

    img_count = 0
    code_count = 0

    for block_type, content in blocks:
        if block_type == 'text':
            process_text_block(doc, content, num_id=num_id)
        elif block_type == 'image':
            add_image_to_doc(doc, content)
            img_count += 1
        elif block_type == 'code':
            para = doc.add_paragraph()
            run = para.add_run(content[:200] + ('...' if len(content) > 200 else ''))
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            code_count += 1

    doc.save(str(docx_path))
    print(f"  Done: {img_count} images embedded, {code_count} code blocks kept")
    return docx_path


def main():
    parser = argparse.ArgumentParser(
        description="Convert SWDD Markdown files to DOCX with embedded PNG images."
    )
    parser.add_argument(
        "--swdd-root", type=Path, default=Path.cwd() / "swdd",
        help="Path to the swdd directory (default: ./swdd)"
    )
    parser.add_argument(
        "modules", nargs="*",
        help="Module names to convert (default: all modules)"
    )
    args = parser.parse_args()

    swdd_dir = args.swdd_root.resolve()
    if not swdd_dir.is_dir():
        print(f"SWDD directory not found: {swdd_dir}")
        sys.exit(1)

    modules = []
    if args.modules:
        for name in args.modules:
            mod_dir = swdd_dir / name
            if mod_dir.is_dir():
                modules.append(mod_dir)
            else:
                print(f"Module directory not found: {mod_dir}")
    else:
        for d in sorted(swdd_dir.iterdir()):
            if d.is_dir() and d.name != "scripts" and (d / "img").exists():
                modules.append(d)

    if not modules:
        print("No modules found to convert.")
        return

    print(f"Converting {len(modules)} modules...\n")
    results = []
    for mod_dir in modules:
        result = convert_module(mod_dir)
        if result:
            results.append(result)
        print()

    print(f"Conversion complete: {len(results)} DOCX files generated.")


if __name__ == "__main__":
    main()
