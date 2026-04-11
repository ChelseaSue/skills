from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
import unicodedata
from copy import copy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from zipfile import BadZipFile

import fitz
import pdfplumber
from docx import Document
from openpyxl import load_workbook


ROOT = Path.cwd()
WORKBOOK_PATH = ROOT / "BBS_相关方需求分析表.xlsx"
ESOW_DIR = ROOT / "ESOW"
REPORT_PATH = ROOT / "stakeholder_requirements_report.json"
BACKUP_PATH = ROOT / f"BBS_相关方需求分析表.backup.{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
SOURCE_SHEET = "相关方输入文档清单"
TARGET_SHEET = "相关方需求分析表"
COL_SEQ = "序号"
COL_NAME = "相关方输入文档名称"
COL_VERSION = "版本"
COL_APPLICABLE = "是否适用"
COL_DISCIPLINE = "分配学科"
COL_REMARK = "备注"
SUPPORTED_EXTS = {".pdf", ".docx", ".doc", ".xls", ".xlsx", ".xlsm", ".docm"}
IGNORE_EXTS = {".dbc", ".ldf"}
FIXED_SATISFY_REASON = "当前按无偏差接收并默认满足，待后续详细分析确认。"
MANUAL_PDF_LINES: dict[str, list[str]] = {
    "osssow": [
        "Open Source Software Instructions for Suppliers",
        "This document provides Volvo Car Corporation and its affiliates’ instructions to Suppliers regarding how to handle Open Source Software related aspects of the contracted assignment.",
        "1. Transparency & Traceability",
        "The Supplier shall provide transparency and traceability in all aspects of OSS compliance related to the contracted assignment, e.g. transparency regarding its own level of OSS maturity and competence, traceability regarding identified OSS compliance issues, and traceability regarding decision making related to OSS compliance, in accordance with the OpenChain Specification 2.1b (ISO/IEC 5230:2020).",
        "2. Training",
        "All key personnel such as OSS Compliance Officer (see 3. Compliance Appointees), software architects and project managers shall have been sufficiently trained prior to the start of the contracted assignment, where sufficiently trained means that key personnel are able to ensure OSS compliant deliverables; see 4. OSS Compliant Deliverables for details.",
        "A summary of the present level of training that the key personnel have undertaken shall be provided at Volvo Cars’ request.",
        "If Volvo Cars deems that the level of training is not sufficient, then Volvo Cars may require that the key personnel attend training arranged by Volvo Cars at Supplier’s expense.",
        "3. OSS Compliance Appointees",
        "The Supplier shall assign one of its employees (hereinafter referred to as the OSS Compliance Officer) to take the overall responsibility for the OSS compliance within the scope of the contracted assignment. The assigned OSS Compliance Officer shall be Volvo Cars’ main contact regarding OSS compliance related queries and issues.",
        "The operative work may be delegated from the OSS Compliance Office to e.g. software architects and project manager(s). However, the overall responsibility remains with the OSS Compliance Officer.",
        "4. OSS Compliant Deliverables",
        "All deliverables shall be 100% OSS compliant, including deliverables such as source code, binaries distributed as separate packages as well as binaries distributed in devices.",
        "1. Identified all code in the deliverable (0% unknown code); see 5. OSS Compliance Reporting for details.",
        "2. Resolved all known OSS compliance issues in the deliverable; see 5. OSS Compliance Reporting for details.",
        "3. Retained attributions in all source code header files; see 6. Coding Rules for details.",
        "4. Provided required attribution in appropriate format, i.e. a License notice file; see 6. Coding Rules for details.",
        "5. Provided all copyleft licensed source code; see 7. Copyleft Archive for details.",
        "Thus, deliverables shall be accompanied by",
        "1. A License notice file",
        "2. A Software Bill of Material",
        "3. A Copyleft Archive",
        "5. OSS Compliance Reporting",
        "The Supplier shall provide an OSS Compliance report regularly. The frequency can be agreed with the appointed Volvo Cars’ OSS Compliance Officer or project manager. However, the minimum, non-negotiable, requirement is that every software delivery has a corresponding OSS Compliance report.",
        "An OSS Compliance report shall be produced by the Supplier using a Volvo Cars approved tool and/or methodology or, upon agreement, the work required to produce a report can be carried out by Volvo Cars on behalf of the Supplier. In the latter case, unless otherwise agreed, the Supplier will carry all Volvo Cars’ expenses for producing reports.",
        "An OSS Compliance report shall contain the following:",
        "1. A Software Bill of Materials (SBOM) listing all software components, including their source path, version number, origin(s) (e.g., “WebKit project”), copyright holder(s) as well as their license(s), preferably with the current version of the SPDX format.",
        "2. A list of known and open OSS compliance issues, including those caused by the deliverable when included in the intended software baseline, and a plan on how and when to resolve them.",
        "3. Ratio of unknown code, including a plan on how and when to bring the ratio of unknown code down to 0%, where unknown code means code which origin(s), copyright holder(s) and/or license(s) are unknown.",
        "6. Coding Rules",
        "The following coding rules shall be followed:",
        "1. A software delivery may only include OSS which is licensed under an open-source license approved by the Open Source Initiative (OSI), unless otherwise agreed in writing by Volvo Cars.",
        "2. Software licensed under AGPL, GPLv3, LGPLv3, or SSPL shall not be included in any deliverable - no exception. A software delivery including AGPL, GPLv3, LGPLv3, or SSPL, intended for inclusion in a device or a service, will be rejected.",
        "3. Consideration regarding work combined with copyleft licensed components, e.g. GPLv2. MPL, EPL etc. must be taken.",
        "4. OSS compliance issues, such as license incompatibilities and incorrect attribution, shall be identified, tracked, and resolved in a timely manner.",
        "5. Always retain (never remove) existing copyright notices from the header of a file, even if only part of the code in the file is used, i.e. copyright attribution shall always be preserved.",
        "6. Copyright attributions, including license text or license permission notices, for all included OSS shall be provided in a collected file, a so called “Notice” file, so it can be presented via the User Interface of the targeted system, product or service, or via the appropriate accompanying documentation. If an application is delivered decoupled from the target system, product or service, the Notice file information shall be presented through application’s own User Interface (e.g. via the application’s About dialog box).",
        "7. All modifications to existing OSS shall be claimed as Volvo Cars’ copyright, unless otherwise agreed in writing by the Parties, by adding “Copyright (c) <YEAR> Volvo Car Corporation”, replacing <YEAR> with the year the modification was done and listed right after existing copyright claims.",
        "8. All modifications to existing OSS which is copyrighted Volvo Cars shall be licensed under a license of Volvo Cars’ choice. Please contact Volvo Cars to receive instructions as to which license (proprietary or open source) shall apply to Volvo Cars’ copyrighted modifications.",
        "9. All essential modifications shall be commented - do follow the standard of the open-source community in question.",
        "10. Volvo Cars’ source code shall never be released publicly, or contributed to an external organisation/community, unless authorised by Volvo Cars.",
        "7. Copyleft Archive",
        "All software deliverables containing components licensed under the so called copyleft licenses, e.g., GPL, LGPL, MPL, EPL, and CDDL, or combined therewith, shall be accompanied by an archive containing all source code of the so licensed components including modifications made.",
        "The source code included in an archive shall be cleaned from all Volvo Cars’ and third parties’ confidential information, such as references to future products etc.",
        "The structure of an archive shall correspond to the structure of the source tree, which means that:",
        "1. separate software modules shall be included directly in that source tree and not as separate sub-archives,",
        "2. deliveries from the Supplier’s suppliers shall be merged into the source tree structure of the Supplier and not provided as separate archives or sub-archives,",
        "3. build instructions, i.e. scripts used to control compilation and installation of the software, must be included.",
        "The archive shall be created as a tar file in bzip2 format (with “tar.bz2” file extension) and shall be created on a computer with a case sensitive file system, e.g. it shall not be created on a Windows computer.",
        "The archive shall be named based on the Volvo Cars label of the corresponding software, e.g., if the Volvo Cars software label is “x.y.z” then the archive shall be named “x.y.z.tar.bz2”.",
        "8. OSS Compliance Audit",
        "Volvo Cars may audit or inspect Supplier or Seller Company’s premises to control and verify the fulfillment of OSS Compliance. Audits or inspections shall take place during normal business hours and with reasonable prior notice, and Supplier shall make available relevant personnel, records, and facilities, as requested by Volvo Cars.",
        "Upon Volvo Cars’ request the Supplier shall make a self-assessment based on a by Volvo Cars provided checklist. The checklist shall be used for the purposes of: i) by Volvo Cars to better understand the OSS related risks with the Supplier, ii) by the Supplier to identify own areas of improvement, and iii) by Volvo Cars and the Supplier in the dialogue regarding continuous improvement of the Supplier’s OSS related capability.",
    ]
}

REQ_PATTERNS = [
    re.compile(r"\bshall\b", re.I),
    re.compile(r"\bmust\b", re.I),
    re.compile(r"\bshould\b", re.I),
    re.compile(r"\bmay\s+not\b", re.I),
    re.compile(r"\bwill\s+not\b", re.I),
    re.compile(r"\bis\s+to\b", re.I),
    re.compile(r"\bare\s+to\b", re.I),
    re.compile(r"\bneeds?\s+to\b", re.I),
    re.compile(r"\brequired\b", re.I),
    re.compile(r"\bcompliant\b", re.I),
    re.compile(r"\bcompatible\b", re.I),
    re.compile(r"\baccording to\b", re.I),
    re.compile(r"必须"),
    re.compile(r"应"),
    re.compile(r"需"),
    re.compile(r"不得"),
    re.compile(r"要求"),
]

ID_PATTERN = re.compile(r"\b([A-Z]{2,}(?:[-_][A-Z0-9]+)+|REQ-\d+)\b")
NUMBERED_PREFIX = re.compile(r"^((?:\d+\.)+(?:\d+)?|(?:\d+\.)|(?:[A-Z]\.)|(?:\d+\)))\s*(.*)$")
SIMPLE_LIST_PREFIX = re.compile(r"^([0-9]+|[A-Z]|[a-z])[\.\)]\s+(.*)$")
META_EXACT = {
    "document name",
    "issuer (dept name, global id number)",
    "issuer",
    "version",
    "date",
    "document type",
    "reg. no.",
    "security class",
    "project:",
    "module name:",
    "buyer dept/name/cds id:",
    "buyer module owner dept/name/cds id:",
    "supplier name:",
    "contractor name:",
    "supplier software quality manager name:",
    "contractor software quality manager name:",
    "accepted | agreed",
    "not accepted | not agreed",
    "accepted with deviation",
    "not applicable",
}


@dataclass
class SourceEntry:
    seq: int
    name: str
    version: str
    discipline: str
    remark: str


@dataclass
class MatchedFile:
    entry: SourceEntry
    path: Path | None
    match_type: str
    detail: str


@dataclass
class RawBlock:
    source_name: str
    version: str
    clause: str
    clause_desc: str
    content: str
    discipline: str
    container: str = ""


@dataclass
class RequirementRow:
    source_name: str
    version: str
    clause: str
    clause_desc: str
    req_type: str
    category: str
    requirement: str
    requirement_id: str
    discipline: str


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    text = text.replace("–", "-").replace("—", "-")
    return text


def normalize_name(name: str) -> str:
    text = normalize_text(name)
    text = re.sub(r"\.[A-Za-z0-9]{1,5}$", "", text)
    text = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "", text.lower())
    return text


def clean_text(text: str) -> str:
    text = normalize_text(text)
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_multiname(raw: str) -> list[str]:
    if not raw:
        return []
    return [part.strip() for part in re.split(r"[\n;；]+", raw) if part.strip()]


def safe_relative(path: Path, root: Path) -> str:
    try:
        return str(path.relative_to(root))
    except Exception:
        return str(path)


def is_ignored_name(name: str) -> bool:
    lower = name.lower().strip()
    return any(lower.endswith(ext) for ext in IGNORE_EXTS)


def read_source_entries() -> list[SourceEntry]:
    wb = load_workbook(WORKBOOK_PATH, data_only=True, read_only=True)
    ws = wb[SOURCE_SHEET]
    headers = [cell or "" for cell in next(ws.iter_rows(min_row=3, max_row=3, values_only=True))]
    col_map = {str(v).strip(): idx for idx, v in enumerate(headers)}
    entries: list[SourceEntry] = []
    for row in ws.iter_rows(min_row=4, values_only=True):
        name_cell = row[col_map[COL_NAME]]
        applicable = str(row[col_map[COL_APPLICABLE]] or "").strip()
        if not name_cell or applicable != "是":
            continue
        if is_ignored_name(str(name_cell)):
            continue
        for name in split_multiname(str(name_cell)):
            if is_ignored_name(name):
                continue
            entries.append(
                SourceEntry(
                    seq=int(row[col_map[COL_SEQ]]),
                    name=name,
                    version=str(row[col_map[COL_VERSION]] or "").strip(),
                    discipline=str(row[col_map[COL_DISCIPLINE]] or "").strip(),
                    remark=str(row[col_map[COL_REMARK]] or "").strip(),
                )
            )
    return entries


def index_esow_files() -> list[Path]:
    paths: list[Path] = []
    for path in ESOW_DIR.rglob("*"):
        if not path.is_file():
            continue
        if path.name.startswith("~$"):
            continue
        suffix = path.suffix.lower()
        if suffix in IGNORE_EXTS:
            continue
        if suffix and suffix not in SUPPORTED_EXTS:
            continue
        paths.append(path)
    return paths


def file_score(entry: SourceEntry, path: Path) -> tuple[int, int, int]:
    score = 0
    if "定点前esow" in path.as_posix().lower():
        score += 30
    if entry.remark and normalize_name(entry.remark) in normalize_name(path.as_posix()):
        score += 20
    suffix = path.suffix.lower()
    if suffix in {".pdf", ".docx", ".doc", ".xlsx", ".xlsm", ".xls"}:
        score += 10
    name_match_len = len(os_common_prefix(normalize_name(entry.name), normalize_name(path.name)))
    return (score, name_match_len, -len(path.as_posix()))


def os_common_prefix(a: str, b: str) -> str:
    max_len = min(len(a), len(b))
    idx = 0
    while idx < max_len and a[idx] == b[idx]:
        idx += 1
    return a[:idx]


def match_files(entries: Iterable[SourceEntry], all_paths: list[Path]) -> list[MatchedFile]:
    exact_index: dict[str, list[Path]] = {}
    norm_index: dict[str, list[Path]] = {}
    for path in all_paths:
        exact_index.setdefault(path.name, []).append(path)
        norm_index.setdefault(normalize_name(path.name), []).append(path)
        norm_index.setdefault(normalize_name(path.stem), []).append(path)

    matched: list[MatchedFile] = []
    seen_paths: set[str] = set()
    for entry in entries:
        candidates: list[tuple[str, list[Path], str]] = []
        if entry.name in exact_index:
            candidates.append(("exact", exact_index[entry.name], entry.name))
        norm = normalize_name(entry.name)
        if norm in norm_index:
            candidates.append(("normalized", norm_index[norm], norm))
        if not candidates:
            contains = [p for p in all_paths if norm and norm in normalize_name(p.name)]
            if contains:
                candidates.append(("contains", contains, norm))

        if not candidates:
            matched.append(MatchedFile(entry, None, "missing", ""))
            continue

        flattened: list[tuple[str, Path, str]] = []
        for match_type, paths, detail in candidates:
            for path in paths:
                flattened.append((match_type, path, detail))
        flattened = list({item[1].as_posix(): item for item in flattened}.values())
        flattened.sort(key=lambda item: file_score(entry, item[1]), reverse=True)
        match_type, best, detail = flattened[0]

        norm_key = normalize_name(entry.name)
        best_key = f"{norm_key}::{best.as_posix()}"
        if norm_key == normalize_name("External testing of material centre EU_21W42") and best_key in seen_paths:
            continue
        seen_paths.add(best_key)
        matched.append(MatchedFile(entry, best, match_type, best.as_posix()))
    return matched


def convert_with_soffice(path: Path, target_ext: str, out_dir: Path) -> Path | None:
    cmd = [
        "libreoffice",
        "--headless",
        "--convert-to",
        target_ext,
        "--outdir",
        str(out_dir),
        str(path),
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=False, timeout=120)
    except subprocess.TimeoutExpired:
        return None
    if result.returncode != 0:
        return None
    converted = out_dir / f"{path.stem}.{target_ext}"
    return converted if converted.exists() else None


def detect_file_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix:
        return suffix
    file_desc = subprocess.run(["file", "-b", str(path)], capture_output=True, text=True, check=False).stdout.lower()
    if "pdf" in file_desc:
        return ".pdf"
    if "word" in file_desc:
        return ".docx"
    if "excel" in file_desc or "spreadsheet" in file_desc:
        return ".xlsx"
    return ""


def extract_pdf_lines(path: Path) -> list[str]:
    lines: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = clean_text(page.extract_text() or "")
                if page_text:
                    lines.extend(page_text.splitlines())
    except Exception:
        lines = []
    meaningful = [ln for ln in lines if re.search(r"[A-Za-z\u4e00-\u9fff]", ln)]
    if len(meaningful) >= 12:
        return merge_wrapped_lines([clean_text(ln) for ln in lines if clean_text(ln)])
    fitz_lines = extract_pdf_lines_fitz(path)
    fitz_meaningful = [ln for ln in fitz_lines if re.search(r"[A-Za-z\u4e00-\u9fff]", ln)]
    if len(fitz_meaningful) >= 12:
        return merge_wrapped_lines(fitz_lines)
    manual = MANUAL_PDF_LINES.get(normalize_name(path.stem))
    if manual:
        return manual
    return merge_wrapped_lines(fitz_lines)


def extract_pdf_lines_fitz(path: Path) -> list[str]:
    doc = fitz.open(path)
    lines: list[str] = []
    for page in doc:
        text = page.get_text("text")
        raw_lines = [clean_text(ln) for ln in text.splitlines() if clean_text(ln)]
        if len([ln for ln in raw_lines if re.search(r"[A-Za-z\u4e00-\u9fff]", ln)]) < 6:
            raw_lines = extract_pdf_lines_fitz_blocks(page)
        lines.extend(raw_lines)
    return lines


def extract_pdf_lines_fitz_blocks(page: fitz.Page) -> list[str]:
    blocks = page.get_text("dict").get("blocks", [])
    collected: list[tuple[float, float, str]] = []
    for block in blocks:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = [span.get("text", "") for span in line.get("spans", [])]
            text = clean_text("".join(spans))
            if text:
                bbox = line.get("bbox", (0, 0, 0, 0))
                collected.append((bbox[1], bbox[0], text))
    collected.sort()
    return [text for _, _, text in collected]


def merge_wrapped_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    for raw in lines:
        line = sanitize_line(raw)
        if not line:
            continue
        if not merged:
            merged.append(line)
            continue
        prev = merged[-1]
        prev_lower = prev.lower()
        line_lower = line.lower()
        if line in {"•", "o", "#"}:
            continue
        if re.fullmatch(r"[0-9]+", line):
            continue
        if prev.endswith(("•", "o", "#", ":", ";", ",")):
            merged[-1] = f"{prev} {line}"
            continue
        if (
            not looks_like_title(prev)
            and not looks_like_title(line)
            and not is_metadata_like(prev)
            and not is_metadata_like(line)
            and not NUMBERED_PREFIX.match(line)
            and not SIMPLE_LIST_PREFIX.match(line)
            and not re.search(r"[.!?。；;:]$", prev)
        ):
            merged[-1] = f"{prev} {line}"
            continue
        if (
            not re.search(r"[.!?。；;:]$", prev)
            and not looks_like_title(line)
            and (
                line[:1].islower()
                or prev_lower.endswith(("the", "a", "an", "of", "to", "with", "and", "or", "for", "in", "on"))
                or len(line.split()) <= 6
            )
        ):
            merged[-1] = f"{prev} {line}"
            continue
        if line.startswith(("• ", "o ", "# ")):
            merged.append(line[2:].strip())
            continue
        merged.append(line)
    return merged


def extract_docx_lines(path: Path) -> list[str]:
    doc = Document(path)
    lines: list[str] = []
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                lines.append(" | ".join(cells))
    return lines


def extract_doc_lines(path: Path) -> list[str]:
    with tempfile.TemporaryDirectory(prefix="doc_convert_") as tmp:
        converted = convert_with_soffice(path, "docx", Path(tmp))
        if not converted:
            return []
        return extract_docx_lines(converted)


def workbook_rows(path: Path) -> list[tuple[str, list[list[str]]]]:
    wb = load_workbook(path, data_only=True, read_only=True)
    out: list[tuple[str, list[list[str]]]] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = [clean_text(str(cell)) for cell in row if cell not in (None, "")]
            cells = [cell for cell in cells if cell]
            if cells:
                rows.append(cells)
        out.append((sheet_name, rows))
    return out


def extract_sheet_rows(path: Path) -> list[tuple[str, list[list[str]]]]:
    try:
        return workbook_rows(path)
    except (BadZipFile, OSError, ValueError):
        with tempfile.TemporaryDirectory(prefix="xlsx_convert_") as tmp:
            converted = convert_with_soffice(path, "xlsx", Path(tmp))
            if not converted:
                return []
            return workbook_rows(converted)


def extract_xls_rows(path: Path) -> list[tuple[str, list[list[str]]]]:
    with tempfile.TemporaryDirectory(prefix="xls_convert_") as tmp:
        converted = convert_with_soffice(path, "xlsx", Path(tmp))
        if not converted:
            return []
        return extract_sheet_rows(converted)


def is_mostly_caps(text: str) -> bool:
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    upper = sum(1 for ch in letters if ch.isupper())
    return upper / len(letters) > 0.75


def looks_like_title(text: str) -> bool:
    raw = clean_text(text)
    if not raw:
        return False
    if len(raw) <= 80 and is_mostly_caps(raw):
        return True
    if (
        len(raw.split()) <= 12
        and raw[:1].isupper()
        and not raw.endswith(".")
        and not any(p.search(raw) for p in REQ_PATTERNS)
    ):
        return True
    if len(raw.split()) <= 10 and not raw.endswith(".") and raw == raw.title():
        return True
    if raw.endswith(":") and len(raw) <= 80:
        return True
    if NUMBERED_PREFIX.match(raw) and len(raw.split()) <= 12:
        return True
    return False


def split_number_prefix(text: str) -> tuple[str, str]:
    raw = clean_text(text)
    match = NUMBERED_PREFIX.match(raw)
    if match:
        return match.group(1).strip(), clean_text(match.group(2))
    return "", raw


def classify_type(text: str, clause: str, clause_desc: str, container: str = "") -> str:
    raw = clean_text(text)
    lower = raw.lower()
    if clause and len(raw.split()) <= 8 and not any(p.search(raw) for p in REQ_PATTERNS):
        return "标题"
    if looks_like_title(raw):
        return "标题"
    if lower.startswith(("note", "purpose", "scope", "description")) and not any(p.search(raw) for p in REQ_PATTERNS):
        return "信息"
    if any(p.search(raw) for p in REQ_PATTERNS):
        return "需求"
    if clause and not any(p.search(raw) for p in REQ_PATTERNS) and len(raw) > 20:
        if clause.count(".") >= 2:
            return "需求"
    return "信息"


def classify_category(text: str, req_type: str, clause_desc: str = "", container: str = "") -> str:
    raw = clean_text(f"{clause_desc} {container} {text}")
    lower = raw.lower()
    if any(token in lower for token in ["test", "audit", "verification", "validate", "inspection", "dvpv"]):
        return "测试需求"
    if any(token in lower for token in ["law", "legal", "regulation", "compliance", "license", "copyright", "osi", "gpl", "lgpl", "agpl", "sspl"]):
        return "法律法规"
    if any(token in lower for token in ["interface", "arxml", "notice file", "sbom", "format", "protocol", "communication", "input/output", "did", "rid"]):
        return "接口需求"
    if any(token in lower for token in ["performance", "latency", "cpu", "rom", "ram", "resource", "periodicity", "timeout", "speed"]):
        return "性能需求"
    if req_type == "标题":
        return "质量需求"
    if any(token in lower for token in ["quality", "traceability", "process", "document", "training", "compatible", "support", "deliverable", "report", "warning", "error", "build environment"]):
        return "质量需求"
    return "功能需求"


def extract_requirement_id(text: str) -> str:
    match = ID_PATTERN.search(text)
    return match.group(1) if match else ""


def is_metadata_like(text: str) -> bool:
    raw = clean_text(text)
    if not raw:
        return True
    lower = raw.lower()
    if lower in META_EXACT:
        return True
    if lower.startswith(("document name", "issuer", "version", "date", "document type", "reg. no.", "security class")):
        return True
    if lower.endswith(":") and len(lower) < 80 and "|" not in lower:
        return True
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}(?: \d{2}:\d{2}:\d{2})?", raw):
        return True
    if re.fullmatch(r"v?\d+(?:\.\d+)+", lower):
        return True
    if re.search(r"\b\d{4}-\d{2}-\d{2}\b", raw):
        return True
    if re.search(r"\bconfidential\b|\bproprietary\b", lower):
        return True
    if re.search(r"\bpage\b|\bversion\b|\bdate\b", lower) and len(raw.split()) <= 12:
        return True
    if re.search(r"\bexternal collaboration operations\b|\bopen source program office\b", lower):
        return True
    return False


def sanitize_line(line: str) -> str:
    line = clean_text(line)
    if not line:
        return ""
    if re.fullmatch(r"[•·\-_*]+", line):
        return ""
    if re.fullmatch(r"\d+\s*(?:\(\d+\))?", line):
        return ""
    if line.lower().startswith(("document name", "issuer (dept", "document type", "security class")):
        return ""
    return line


def parse_text_lines(lines: list[str], entry: SourceEntry, source_name: str) -> list[RawBlock]:
    results: list[RawBlock] = []
    current_heading = source_name
    current_clause = ""
    in_list_section = ""
    current_list_items: list[str] = []

    def flush_list_items() -> None:
        nonlocal current_list_items
        if not current_list_items:
            return
        for item in current_list_items:
            prefix, body = split_number_prefix(item)
            clause = f"{in_list_section} / {prefix}" if in_list_section and prefix else current_clause
            content = body or item
            results.append(
                RawBlock(
                    source_name=entry.name,
                    version=entry.version,
                    clause=clause.strip(" /"),
                    clause_desc=current_heading,
                    content=content,
                    discipline=entry.discipline,
                    container=current_heading,
                )
            )
        current_list_items = []

    for raw in lines:
        line = sanitize_line(raw)
        if not line:
            continue
        clause_prefix, body = split_number_prefix(line)
        if clause_prefix and body and looks_like_title(line):
            flush_list_items()
            current_clause = clause_prefix
            current_heading = body
            in_list_section = f"{clause_prefix} {body}".strip()
            results.append(
                RawBlock(
                    source_name=entry.name,
                    version=entry.version,
                    clause=clause_prefix,
                    clause_desc=body,
                    content=body,
                    discipline=entry.discipline,
                    container=source_name,
                )
            )
            continue
        if looks_like_title(line):
            flush_list_items()
            current_heading = line.rstrip(":")
            if clause_prefix:
                current_clause = clause_prefix
            results.append(
                RawBlock(
                    source_name=entry.name,
                    version=entry.version,
                    clause=current_clause,
                    clause_desc=current_heading,
                    content=current_heading,
                    discipline=entry.discipline,
                    container=source_name,
                )
            )
            continue
        if in_list_section and SIMPLE_LIST_PREFIX.match(line):
            current_list_items.append(line)
            continue
        flush_list_items()
        results.append(
            RawBlock(
                source_name=entry.name,
                version=entry.version,
                clause=clause_prefix or current_clause,
                clause_desc=current_heading,
                content=body if clause_prefix and body else line,
                discipline=entry.discipline,
                container=source_name,
            )
        )
    flush_list_items()
    return results


def parse_sheet_rows(rows_by_sheet: list[tuple[str, list[list[str]]]], entry: SourceEntry, source_name: str) -> list[RawBlock]:
    results: list[RawBlock] = []
    for sheet_name, rows in rows_by_sheet:
        current_heading = sheet_name
        results.append(
            RawBlock(
                source_name=entry.name,
                version=entry.version,
                clause=sheet_name,
                clause_desc=sheet_name,
                content=sheet_name,
                discipline=entry.discipline,
                container=sheet_name,
            )
        )
        for cells in rows:
            joined = " | ".join(cells)
            if not joined:
                continue
            first = cells[0]
            second = cells[1] if len(cells) > 1 else ""
            if first.lower() == "ref." and "requirement" in joined.lower():
                continue
            if len(cells) == 1 and looks_like_title(first):
                current_heading = first
                clause, body = split_number_prefix(first)
                results.append(
                    RawBlock(
                        source_name=entry.name,
                        version=entry.version,
                        clause=clause or sheet_name,
                        clause_desc=body or first,
                        content=body or first,
                        discipline=entry.discipline,
                        container=sheet_name,
                    )
                )
                continue
            if NUMBERED_PREFIX.match(first):
                clause, maybe_title = split_number_prefix(first)
                content = second or maybe_title or joined
                if len(cells) == 2 and maybe_title and not second:
                    current_heading = maybe_title
                elif len(cells) == 2 and maybe_title and looks_like_title(maybe_title):
                    current_heading = maybe_title
                results.append(
                    RawBlock(
                        source_name=entry.name,
                        version=entry.version,
                        clause=clause,
                        clause_desc=current_heading,
                        content=content,
                        discipline=entry.discipline,
                        container=sheet_name,
                    )
                )
                continue
            if first.lower().startswith("note"):
                results.append(
                    RawBlock(
                        source_name=entry.name,
                        version=entry.version,
                        clause="",
                        clause_desc=current_heading,
                        content=joined,
                        discipline=entry.discipline,
                        container=sheet_name,
                    )
                )
                continue
            if len(cells) <= 3 and all(len(cell.split()) <= 8 for cell in cells):
                title = joined
                current_heading = title
                results.append(
                    RawBlock(
                        source_name=entry.name,
                        version=entry.version,
                        clause="",
                        clause_desc=current_heading,
                        content=title,
                        discipline=entry.discipline,
                        container=sheet_name,
                    )
                )
                continue
            results.append(
                RawBlock(
                    source_name=entry.name,
                    version=entry.version,
                    clause="",
                    clause_desc=current_heading,
                    content=joined,
                    discipline=entry.discipline,
                    container=sheet_name,
                )
            )
    return results


def extract_raw_blocks(path: Path, entry: SourceEntry) -> list[RawBlock]:
    file_type = detect_file_type(path)
    if file_type == ".pdf":
        lines = extract_pdf_lines(path)
        return parse_text_lines(lines, entry, path.stem)
    if file_type in {".docx", ".docm"}:
        lines = extract_docx_lines(path)
        return parse_text_lines(lines, entry, path.stem)
    if file_type == ".doc":
        lines = extract_doc_lines(path)
        return parse_text_lines(lines, entry, path.stem)
    if file_type in {".xlsx", ".xlsm"}:
        rows = extract_sheet_rows(path)
        return parse_sheet_rows(rows, entry, path.stem)
    if file_type == ".xls":
        rows = extract_xls_rows(path)
        return parse_sheet_rows(rows, entry, path.stem)
    return []


def raw_block_to_row(block: RawBlock) -> RequirementRow | None:
    content = clean_text(block.content)
    if not content:
        return None
    if is_metadata_like(content):
        return None
    req_type = classify_type(content, block.clause, block.clause_desc, block.container)
    category = classify_category(content, req_type, block.clause_desc, block.container) if req_type == "需求" else ""
    requirement_id = extract_requirement_id(content)[:200] if req_type == "需求" else ""
    return RequirementRow(
        source_name=block.source_name,
        version=block.version,
        clause=block.clause[:200],
        clause_desc=(block.clause_desc or block.container or block.source_name)[:500],
        req_type=req_type,
        category=category,
        requirement=content[:12000],
        requirement_id=requirement_id,
        discipline=block.discipline,
    )


def dedupe_rows(rows: list[RequirementRow]) -> list[RequirementRow]:
    deduped: list[RequirementRow] = []
    seen: set[tuple[str, str, str, str]] = set()
    for row in rows:
        key = (
            normalize_name(row.source_name),
            normalize_name(row.clause),
            normalize_name(row.requirement),
            row.req_type,
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def looks_like_continuation(prev: str, curr: str) -> bool:
    prev = clean_text(prev)
    curr = clean_text(curr)
    if not prev or not curr:
        return False
    if curr[:1].islower():
        return True
    if prev.endswith((":", ";", ",", "for each", "for", "to", "with", "and", "or")):
        return True
    if not re.search(r"[.!?。；;:]$", prev):
        return True
    return False


def merge_adjacent_rows(rows: list[RequirementRow]) -> list[RequirementRow]:
    if not rows:
        return []
    merged: list[RequirementRow] = [rows[0]]
    for row in rows[1:]:
        prev = merged[-1]
        if (
            prev.source_name == row.source_name
            and prev.version == row.version
            and prev.clause == row.clause
            and prev.clause_desc == row.clause_desc
            and prev.req_type != "标题"
            and row.req_type != "标题"
            and looks_like_continuation(prev.requirement, row.requirement)
        ):
            combined = f"{prev.requirement.rstrip()} {row.requirement.lstrip()}"
            req_type = "需求" if "需求" in {prev.req_type, row.req_type} else "信息"
            merged[-1] = RequirementRow(
                source_name=prev.source_name,
                version=prev.version,
                clause=prev.clause,
                clause_desc=prev.clause_desc,
                req_type=req_type,
                category=classify_category(combined, req_type, prev.clause_desc, "") if req_type == "需求" else "",
                requirement=combined[:12000],
                requirement_id=extract_requirement_id(combined)[:200] if req_type == "需求" else "",
                discipline=prev.discipline,
            )
            continue
        merged.append(row)
    return merged


def capture_row_style(ws, template_row: int, start_col: int = 1, end_col: int = 26) -> dict[str, object]:
    style_map: dict[int, dict[str, object]] = {}
    for col in range(start_col, end_col + 1):
        src = ws.cell(template_row, col)
        style_map[col] = {
            "has_style": src.has_style,
            "font": copy(src.font),
            "fill": copy(src.fill),
            "border": copy(src.border),
            "alignment": copy(src.alignment),
            "number_format": src.number_format,
            "protection": copy(src.protection),
        }
    return {
        "height": ws.row_dimensions[template_row].height,
        "styles": style_map,
    }


def apply_row_style(ws, cached_style: dict[str, object], target_row: int, start_col: int = 1, end_col: int = 26) -> None:
    ws.row_dimensions[target_row].height = cached_style["height"]
    style_map = cached_style["styles"]
    for col in range(start_col, end_col + 1):
        dst = ws.cell(target_row, col)
        src_style = style_map[col]
        if src_style["has_style"]:
            dst.font = copy(src_style["font"])
            dst.fill = copy(src_style["fill"])
            dst.border = copy(src_style["border"])
            dst.alignment = copy(src_style["alignment"])
            dst.number_format = src_style["number_format"]
            dst.protection = copy(src_style["protection"])


def write_rows(rows: list[RequirementRow]) -> None:
    shutil.copy2(WORKBOOK_PATH, BACKUP_PATH)
    wb = load_workbook(WORKBOOK_PATH)
    ws = wb[TARGET_SHEET]
    template_values = [ws.cell(4, col).value for col in range(1, 27)]
    template_style = capture_row_style(ws, 4)
    if ws.max_row >= 4:
        ws.delete_rows(4, ws.max_row - 3)

    for idx, item in enumerate(rows, start=4):
        apply_row_style(ws, template_style, idx)
        for col, value in enumerate(template_values, start=1):
            ws.cell(idx, col).value = value if col == 1 else None
        seq = idx - 3
        payload = {
            2: seq,
            3: item.source_name,
            4: item.version,
            5: item.clause,
            6: item.clause_desc,
            7: item.req_type,
            8: item.category,
            9: "是",
            10: "",
            11: "否",
            12: "",
            13: "",
            14: "",
            15: item.requirement,
            16: item.requirement_id,
            17: "是",
            18: FIXED_SATISFY_REASON,
            19: "",
            20: "",
            21: "",
            22: "",
            23: "",
            24: item.discipline,
            25: "已接受",
            26: "",
        }
        for col, value in payload.items():
            ws.cell(idx, col).value = value
    wb.save(WORKBOOK_PATH)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract stakeholder requirement items from source documents and fill an Excel template.")
    parser.add_argument("--workbook", required=True, help="Path to the target Excel workbook.")
    parser.add_argument("--source-root", required=True, help="Root directory containing source documents.")
    parser.add_argument("--report-path", help="Optional output path for the JSON report. Defaults beside workbook.")
    parser.add_argument("--source-sheet", default="相关方输入文档清单", help="Sheet name containing the input document list.")
    parser.add_argument("--target-sheet", default="相关方需求分析表", help="Sheet name to write extracted rows into.")
    parser.add_argument("--col-seq", default="序号", help="Header name for sequence column in source sheet.")
    parser.add_argument("--col-name", default="相关方输入文档名称", help="Header name for source document name column.")
    parser.add_argument("--col-version", default="版本", help="Header name for version column.")
    parser.add_argument("--col-applicable", default="是否适用", help="Header name for applicability column.")
    parser.add_argument("--col-discipline", default="分配学科", help="Header name for discipline column.")
    parser.add_argument("--col-remark", default="备注", help="Header name for remark column.")
    return parser.parse_args()


def configure_runtime(args: argparse.Namespace) -> None:
    global ROOT, WORKBOOK_PATH, ESOW_DIR, REPORT_PATH, BACKUP_PATH
    global SOURCE_SHEET, TARGET_SHEET, COL_SEQ, COL_NAME, COL_VERSION, COL_APPLICABLE, COL_DISCIPLINE, COL_REMARK

    WORKBOOK_PATH = Path(args.workbook).resolve()
    ESOW_DIR = Path(args.source_root).resolve()
    ROOT = WORKBOOK_PATH.parent
    REPORT_PATH = Path(args.report_path).resolve() if args.report_path else ROOT / "stakeholder_requirements_report.json"
    BACKUP_PATH = ROOT / f"{WORKBOOK_PATH.stem}.backup.{datetime.now().strftime('%Y%m%d_%H%M%S')}{WORKBOOK_PATH.suffix}"

    SOURCE_SHEET = args.source_sheet
    TARGET_SHEET = args.target_sheet
    COL_SEQ = args.col_seq
    COL_NAME = args.col_name
    COL_VERSION = args.col_version
    COL_APPLICABLE = args.col_applicable
    COL_DISCIPLINE = args.col_discipline
    COL_REMARK = args.col_remark


def main() -> None:
    args = parse_args()
    configure_runtime(args)
    entries = read_source_entries()
    all_paths = index_esow_files()
    matched = match_files(entries, all_paths)

    report: dict[str, object] = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "workbook": WORKBOOK_PATH.name,
        "backup": BACKUP_PATH.name,
        "stats": {},
        "matches": [],
        "missing": [],
        "parse_failures": [],
        "ignored": sorted(IGNORE_EXTS),
    }

    requirement_rows: list[RequirementRow] = []
    for idx, item in enumerate(matched, start=1):
        print(f"[{idx}/{len(matched)}] {item.entry.name}", flush=True)
        match_info = {
            "name": item.entry.name,
            "version": item.entry.version,
            "match_type": item.match_type,
            "detail": item.detail,
            "path": safe_relative(item.path, ROOT) if item.path else "",
        }
        if not item.path:
            report["missing"].append(match_info)
            continue
        report["matches"].append(match_info)
        try:
            raw_blocks = extract_raw_blocks(item.path, item.entry)
        except Exception as exc:
            report["parse_failures"].append({**match_info, "reason": f"extract_error:{type(exc).__name__}"})
            continue
        if not raw_blocks:
            report["parse_failures"].append({**match_info, "reason": "no_content_extracted"})
            continue
        added = 0
        for block in raw_blocks:
            row = raw_block_to_row(block)
            if not row:
                continue
            requirement_rows.append(row)
            added += 1
        if added == 0:
            report["parse_failures"].append({**match_info, "reason": "no_row_generated"})

    final_rows = merge_adjacent_rows(dedupe_rows(requirement_rows))
    write_rows(final_rows)

    report["stats"] = {
        "applicable_entries": len(entries),
        "matched_files": len(report["matches"]),
        "missing_files": len(report["missing"]),
        "parse_failures": len(report["parse_failures"]),
        "written_requirements": len(final_rows),
    }
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report["stats"], ensure_ascii=False, indent=2))
    print(f"REPORT: {REPORT_PATH}")
    print(f"BACKUP: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
